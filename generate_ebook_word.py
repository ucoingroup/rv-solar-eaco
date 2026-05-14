#!/usr/bin/env python3
"""
EACO WEB3 WEB4 Guide - Word Document Generator
Generates a complete Word document from Markdown chapter files
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown

# Configuration
WORKSPACE = r"C:\Users\Administrator\.qclaw\workspace-ua58rsb93veqtxl7"
CHAPTER_FILES = [
    "EACO_WEB3_WEB4_Guide_Chapters_001_010.md",
    "EACO_WEB3_WEB4_Guide_Chapters_011_020.md",
    "EACO_WEB3_WEB4_Guide_Chapters_021_030.md",
    "EACO_WEB3_WEB4_Guide_Chapters_031_040.md",
    "EACO_WEB3_WEB4_Guide_Chapters_041_050.md",
    "EACO_WEB3_WEB4_Guide_Chapters_051_070.md",
    "EACO_WEB3_WEB4_Guide_Chapters_071_100.md",
]
OUTPUT_FILE = "EACO_WEB3_WEB4_Guide_Complete.docx"

def read_markdown_file(filepath):
    """Read markdown file content"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_heading(doc, text, level):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para

def add_table(doc, content):
    """Parse and add a table from markdown format"""
    lines = content.strip().split('\n')
    
    # Find table rows (skip header separator)
    rows_data = []
    for line in lines:
        line = line.strip()
        if line and '|' in line and not line.replace('|', '').replace('-', '').replace(':', '').strip() == '':
            if not re.match(r'^\|[\s\-:|]+\|$', line):  # Skip separator lines
                # Parse cells
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]  # Remove empty
                if cells:
                    rows_data.append(cells)
    
    if rows_data:
        # Determine max columns
        max_cols = max(len(row) for row in rows_data)
        min_cols = min(len(row) for row in rows_data)
        
        # If inconsistent columns, normalize to max
        if min_cols != max_cols:
            for row in rows_data:
                while len(row) < max_cols:
                    row.append('')
        
        try:
            # Create table
            table = doc.add_table(rows=len(rows_data), cols=max_cols)
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(rows_data):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = cell_text
                        # Bold first column if it looks like a header
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            return True
        except Exception as e:
            print(f"Warning: Table parsing error: {e}")
            return False
    return False

def process_markdown_to_docx(doc, md_content):
    """Convert markdown content to docx paragraphs"""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_content = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip title lines at the beginning (book title)
        if i < 5 and line.startswith('# '):
            i += 1
            continue
        
        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_heading(doc, text, 1)
            i += 1
            continue
        
        # Heading 2
        elif line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            add_heading(doc, text, 2)
            i += 1
            continue
        
        # Heading 3
        elif line.startswith('### '):
            text = line[4:].strip()
            add_heading(doc, text, 3)
            i += 1
            continue
        
        # Heading 4
        elif line.startswith('#### '):
            text = line[5:].strip()
            add_heading(doc, text, 4)
            i += 1
            continue
        
        # Table detection
        elif '|' in line and line.strip().startswith('|'):
            # Start collecting table rows
            table_content = []
            while i < len(lines) and '|' in lines[i]:
                table_content.append(lines[i])
                i += 1
            add_table(doc, '\n'.join(table_content))
            continue
        
        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            # Add code block
            para = doc.add_paragraph()
            run = para.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            para.paragraph_format.left_indent = Inches(0.3)
            continue
        
        # Unordered list
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:].strip()
                items.append(item_text)
                i += 1
            for item in items:
                para = doc.add_paragraph(item, style='List Bullet')
            continue
        
        # Ordered list
        elif re.match(r'^\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                item_text = re.sub(r'^\d+\.\s', '', lines[i]).strip()
                items.append(item_text)
                i += 1
            for item in items:
                para = doc.add_paragraph(item, style='List Number')
            continue
        
        # Empty line
        elif line.strip() == '':
            i += 1
            continue
        
        # Regular paragraph
        else:
            # Clean markdown formatting
            text = line
            
            # Remove bold markers but keep text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'__(.+?)__', r'\1', text)
            
            # Remove italic markers but keep text
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'_(.+?)_', r'\1', text)
            
            # Remove inline code markers
            text = re.sub(r'`(.+?)`', r'\1', text)
            
            # Remove links but keep text
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            
            # Clean up extra spaces
            text = re.sub(r'\s+', ' ', text)
            
            if text.strip():
                add_paragraph(doc, text)
            i += 1
            continue
        
        i += 1

def main():
    """Main function to generate the complete Word document"""
    print("Starting EACO WEB3 WEB4 Guide Word Document Generator...")
    
    # Create new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "WEB3 & WEB4 智能科技终极指南 2026"
    doc.core_properties.author = "EACO Community"
    doc.core_properties.subject = "EACO - Earth's Best Coin"
    
    # Add title page
    title = doc.add_heading('WEB3 & WEB4 智能科技终极指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('EACO - Earth\'s Best Coin')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 100, 200)
    
    year_para = doc.add_paragraph('2026 Edition')
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_para.runs[0].font.size = Pt(14)
    
    doc.add_page_break()
    
    # Process each chapter file
    total_chapters = 0
    for filename in CHAPTER_FILES:
        filepath = os.path.join(WORKSPACE, filename)
        if os.path.exists(filepath):
            print(f"Processing: {filename}")
            try:
                content = read_markdown_file(filepath)
                process_markdown_to_docx(doc, content)
                total_chapters += 1
                # Add page break between files
                doc.add_page_break()
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                continue
        else:
            print(f"Warning: File not found - {filepath}")
    
    # Save document
    output_path = os.path.join(WORKSPACE, OUTPUT_FILE)
    doc.save(output_path)
    print(f"\n✅ Success! Word document generated: {output_path}")
    print(f"   Processed {total_chapters} chapter files")
    print(f"   Total chapters: 1-100")
    
    return output_path

if __name__ == "__main__":
    main()
